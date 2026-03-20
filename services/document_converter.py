import base64
import io
import json
import os
import shutil
import uuid
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from threading import Lock
from typing import Any, Dict, List, Optional, Tuple

import fitz
import pytesseract
from docx import Document
from PIL import Image
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename

try:
    from openai import OpenAI
except Exception:
    OpenAI = None


class DocumentConverterService:
    allowed_formats = {"pdf", "docx", "txt", "png", "jpg", "jpeg"}
    conversion_map = {
        "pdf": {"docx", "txt", "png"},
        "docx": {"pdf", "txt"},
        "txt": {"docx", "pdf"},
        "png": {"pdf", "txt"},
        "jpg": {"pdf", "txt"},
        "jpeg": {"pdf", "txt"},
    }

    def __init__(self, base_dir: Path, store):
        self.base_dir = base_dir
        self.store = store
        self.data_dir = base_dir / "data" / "converter"
        self.uploads_dir = self.data_dir / "uploads"
        self.outputs_dir = self.data_dir / "outputs"
        self.previews_dir = self.data_dir / "previews"
        self._lock = Lock()
        for folder in (self.uploads_dir, self.outputs_dir, self.previews_dir):
            folder.mkdir(parents=True, exist_ok=True)

    def _provider_client(self) -> Tuple[Optional[Any], str]:
        provider = os.getenv("OCR_PROVIDER", os.getenv("GENERATOR_PROVIDER", "groq")).strip().lower()
        if OpenAI is None:
            return None, ""
        if provider == "groq":
            key = os.getenv("OCR_API_KEY", os.getenv("GROQ_API_KEY", "")).strip()
            base_url = os.getenv("OCR_BASE_URL", os.getenv("GROQ_BASE_URL", "https://api.groq.com/openai/v1")).strip()
            model = os.getenv("OCR_MODEL", os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")).strip()
        else:
            key = os.getenv("OCR_API_KEY", os.getenv("OPENAI_API_KEY", "")).strip()
            base_url = os.getenv("OCR_BASE_URL", os.getenv("OPENAI_BASE_URL", "")).strip()
            model = os.getenv("OCR_MODEL", os.getenv("OPENAI_MODEL", "gpt-4.1-mini")).strip()
        if not key:
            return None, model
        kwargs: Dict[str, Any] = {"api_key": key}
        if base_url:
            kwargs["base_url"] = base_url
        return OpenAI(**kwargs), model

    def list_jobs(self, owner_uid: str) -> List[Dict[str, Any]]:
        return self.store.list_converter_jobs(owner_uid)

    def list_all_jobs(self, limit: int = 200) -> List[Dict[str, Any]]:
        return self.store.list_all_converter_jobs(limit)

    def get_job(self, job_id: str) -> Optional[Dict[str, Any]]:
        return self.store.get_converter_job(job_id)

    def get_upload(self, upload_id: str) -> Optional[Dict[str, Any]]:
        return self.store.get_converter_upload(upload_id)

    def list_uploads(self, owner_uid: str) -> List[Dict[str, Any]]:
        return self.store.list_converter_uploads(owner_uid)

    def save_uploads(self, files: List[FileStorage], owner_uid: str, owner_email: str) -> List[Dict[str, Any]]:
        saved: List[Dict[str, Any]] = []
        for file in files:
            original_name = secure_filename(file.filename or "")
            ext = original_name.rsplit(".", 1)[-1].lower() if "." in original_name else ""
            if not original_name or ext not in self.allowed_formats:
                continue
            upload_id = uuid.uuid4().hex[:12]
            stored_name = f"{upload_id}-{original_name}"
            output_path = self.uploads_dir / stored_name
            file.save(output_path)
            preview = self._build_source_preview(output_path, ext)
            record = {
                "id": upload_id,
                "owner_uid": owner_uid,
                "owner_email": owner_email,
                "fileName": original_name,
                "storedPath": str(output_path),
                "sourceFormat": ext,
                "sizeBytes": output_path.stat().st_size,
                "createdAt": datetime.now(timezone.utc).isoformat(),
                "preview": preview,
            }
            self.store.create_converter_upload(record)
            saved.append(record)
        return saved

    def convert(
        self,
        upload_id: str,
        owner_uid: str,
        owner_email: str,
        target_format: str,
        ai_mode: str,
        ocr_enabled: bool,
        structure_fix: bool,
        keep_layout: bool,
        priority: str,
    ) -> Dict[str, Any]:
        upload = self.get_upload(upload_id)
        if not upload:
            raise ValueError("Upload not found.")
        if upload.get("owner_uid") != owner_uid:
            raise ValueError("You do not have access to this uploaded file.")

        source_format = upload["sourceFormat"]
        target_format = target_format.lower().strip()
        if target_format not in self.conversion_map.get(source_format, set()):
            raise ValueError(f"Unsupported conversion: {source_format} -> {target_format}")

        input_path = Path(upload["storedPath"])
        job_id = uuid.uuid4().hex[:12]
        result = self._run_conversion(
            input_path=input_path,
            source_format=source_format,
            target_format=target_format,
            ai_mode=ai_mode,
            ocr_enabled=ocr_enabled,
            structure_fix=structure_fix,
            keep_layout=keep_layout,
            job_id=job_id,
        )

        job = {
            "id": job_id,
            "uploadId": upload_id,
            "owner_uid": owner_uid,
            "owner_email": owner_email,
            "fileName": upload["fileName"],
            "sourceFormat": source_format,
            "targetFormat": target_format,
            "aiMode": ai_mode,
            "priority": priority,
            "ocrEnabled": ocr_enabled,
            "structureFix": structure_fix,
            "keepLayout": keep_layout,
            "status": "completed",
            "createdAt": datetime.now(timezone.utc).isoformat(),
            "downloadName": result["downloadName"],
            "downloadPath": result["downloadPath"],
            "preview": result["preview"],
            "notes": result.get("notes", []),
        }
        self.store.create_converter_job(job)
        return job

    def cleanup_jobs(self, max_age_days: int, owner_uid: Optional[str] = None) -> Dict[str, int]:
        now = datetime.now(timezone.utc)
        candidates = self.list_jobs(owner_uid) if owner_uid else self.list_all_jobs(1000)
        deleted_jobs = 0
        deleted_uploads = 0
        seen_uploads = set()
        for job in candidates:
            created_at = job.get("createdAt", "")
            try:
                created = datetime.fromisoformat(created_at.replace("Z", "+00:00"))
            except Exception:
                continue
            if (now - created).days < max_age_days:
                continue
            download_path = Path(job.get("downloadPath") or "")
            if download_path.exists():
                if download_path.is_file():
                    download_path.unlink(missing_ok=True)
                elif download_path.is_dir():
                    shutil.rmtree(download_path, ignore_errors=True)
            self.store.delete_converter_job(job["id"])
            deleted_jobs += 1
            upload_id = job.get("uploadId") or ""
            if upload_id and upload_id not in seen_uploads:
                upload = self.get_upload(upload_id)
                if upload:
                    upload_path = Path(upload.get("storedPath") or "")
                    if upload_path.exists():
                        upload_path.unlink(missing_ok=True)
                    self.store.delete_converter_upload(upload_id)
                    deleted_uploads += 1
                seen_uploads.add(upload_id)
        return {"deleted_jobs": deleted_jobs, "deleted_uploads": deleted_uploads}

    def delete_job(self, job_id: str, owner_uid: str, is_admin: bool = False) -> bool:
        job = self.get_job(job_id)
        if not job:
            return False
        if not is_admin and job.get("owner_uid") != owner_uid:
            return False
        download_path = Path(job.get("downloadPath") or "")
        if download_path.exists():
            if download_path.is_file():
                download_path.unlink(missing_ok=True)
            elif download_path.is_dir():
                shutil.rmtree(download_path, ignore_errors=True)
        sibling_dir = self.outputs_dir / f"{job_id}-pages"
        if sibling_dir.exists():
            shutil.rmtree(sibling_dir, ignore_errors=True)
        upload_id = job.get("uploadId") or ""
        if upload_id:
            upload = self.get_upload(upload_id)
            if upload and (is_admin or upload.get("owner_uid") == owner_uid):
                upload_path = Path(upload.get("storedPath") or "")
                if upload_path.exists():
                    upload_path.unlink(missing_ok=True)
                self.store.delete_converter_upload(upload_id)
        return self.store.delete_converter_job(job_id)

    def _run_conversion(
        self,
        input_path: Path,
        source_format: str,
        target_format: str,
        ai_mode: str,
        ocr_enabled: bool,
        structure_fix: bool,
        keep_layout: bool,
        job_id: str,
    ) -> Dict[str, Any]:
        notes: List[str] = []
        if source_format == "docx" and target_format == "txt":
            text = self._docx_to_text(input_path)
            text = self._cleanup_text(text, ai_mode, structure_fix)
            output_path = self.outputs_dir / f"{job_id}.txt"
            output_path.write_text(text, encoding="utf-8")
            preview = {"kind": "text", "text": text[:6000]}
            return {"downloadName": output_path.name, "downloadPath": str(output_path), "preview": preview, "notes": notes}

        if source_format == "txt" and target_format == "docx":
            text = input_path.read_text(encoding="utf-8", errors="ignore")
            text = self._cleanup_text(text, ai_mode, structure_fix)
            output_path = self.outputs_dir / f"{job_id}.docx"
            self._text_to_docx(text, output_path)
            preview = {"kind": "text", "text": text[:6000]}
            return {"downloadName": output_path.name, "downloadPath": str(output_path), "preview": preview, "notes": notes}

        if source_format == "docx" and target_format == "pdf":
            text = self._docx_to_text(input_path)
            text = self._cleanup_text(text, ai_mode, structure_fix)
            output_path = self.outputs_dir / f"{job_id}.pdf"
            self._text_to_pdf(text, output_path)
            preview = {"kind": "text", "text": text[:6000]}
            return {"downloadName": output_path.name, "downloadPath": str(output_path), "preview": preview, "notes": notes}

        if source_format == "pdf" and target_format == "docx":
            text, ocr_notes = self._pdf_to_text(input_path, ai_mode, ocr_enabled, structure_fix)
            notes.extend(ocr_notes)
            output_path = self.outputs_dir / f"{job_id}.docx"
            self._text_to_docx(text, output_path)
            preview = {"kind": "text", "text": text[:6000]}
            return {"downloadName": output_path.name, "downloadPath": str(output_path), "preview": preview, "notes": notes}

        if source_format == "pdf" and target_format == "txt":
            text, ocr_notes = self._pdf_to_text(input_path, ai_mode, ocr_enabled, structure_fix)
            notes.extend(ocr_notes)
            output_path = self.outputs_dir / f"{job_id}.txt"
            output_path.write_text(text, encoding="utf-8")
            preview = {"kind": "text", "text": text[:6000]}
            return {"downloadName": output_path.name, "downloadPath": str(output_path), "preview": preview, "notes": notes}

        if source_format == "pdf" and target_format == "png":
            bundle_path, preview = self._pdf_to_images(input_path, job_id)
            if bundle_path.suffix.lower() == ".zip":
                notes.append("Multi-page PDF exported as ZIP of PNG images.")
            return {"downloadName": bundle_path.name, "downloadPath": str(bundle_path), "preview": preview, "notes": notes}

        if source_format in {"png", "jpg", "jpeg"} and target_format == "pdf":
            output_path = self.outputs_dir / f"{job_id}.pdf"
            self._image_to_pdf(input_path, output_path)
            preview = {"kind": "image", "url": f"/api/converter/assets/{output_path.name}"}
            return {"downloadName": output_path.name, "downloadPath": str(output_path), "preview": preview, "notes": notes}

        if source_format in {"png", "jpg", "jpeg"} and target_format == "txt":
            text = self._image_to_text(input_path, ai_mode, ocr_enabled, structure_fix, notes)
            output_path = self.outputs_dir / f"{job_id}.txt"
            output_path.write_text(text, encoding="utf-8")
            preview = {"kind": "text", "text": text[:6000]}
            return {"downloadName": output_path.name, "downloadPath": str(output_path), "preview": preview, "notes": notes}

        if source_format == "txt" and target_format == "pdf":
            text = input_path.read_text(encoding="utf-8", errors="ignore")
            text = self._cleanup_text(text, ai_mode, structure_fix)
            output_path = self.outputs_dir / f"{job_id}.pdf"
            self._text_to_pdf(text, output_path)
            preview = {"kind": "text", "text": text[:6000]}
            return {"downloadName": output_path.name, "downloadPath": str(output_path), "preview": preview, "notes": notes}

        raise ValueError("Conversion path is not implemented.")

    def _docx_to_text(self, input_path: Path) -> str:
        document = Document(str(input_path))
        lines = [paragraph.text.strip() for paragraph in document.paragraphs]
        return "\n".join(line for line in lines if line).strip()

    def _text_to_docx(self, text: str, output_path: Path) -> None:
        document = Document()
        for block in text.splitlines():
            if block.strip():
                document.add_paragraph(block.strip())
            else:
                document.add_paragraph("")
        document.save(str(output_path))

    def _text_to_pdf(self, text: str, output_path: Path) -> None:
        packet = canvas.Canvas(str(output_path), pagesize=A4)
        width, height = A4
        y = height - 48
        text_obj = packet.beginText(48, y)
        text_obj.setFont("Helvetica", 11)
        for raw_line in text.splitlines() or [""]:
            line = raw_line[:120]
            text_obj.textLine(line)
            if text_obj.getY() <= 48:
                packet.drawText(text_obj)
                packet.showPage()
                text_obj = packet.beginText(48, height - 48)
                text_obj.setFont("Helvetica", 11)
        packet.drawText(text_obj)
        packet.save()

    def _pdf_to_text(self, input_path: Path, ai_mode: str, ocr_enabled: bool, structure_fix: bool) -> Tuple[str, List[str]]:
        notes: List[str] = []
        doc = fitz.open(str(input_path))
        pages: List[str] = []
        for idx, page in enumerate(doc):
            extracted = page.get_text("text").strip()
            if extracted:
                pages.append(extracted)
                continue
            if not ocr_enabled:
                pages.append("")
                notes.append(f"Page {idx + 1}: no embedded text and OCR disabled.")
                continue
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image_bytes = pix.tobytes("png")
            ocr_text, mode_note = self._ocr_image_bytes(image_bytes, ai_mode)
            if mode_note:
                notes.append(f"Page {idx + 1}: {mode_note}")
            pages.append(ocr_text)
        full_text = "\n\n".join(page for page in pages if page).strip()
        return self._cleanup_text(full_text, ai_mode, structure_fix), notes

    def _pdf_to_images(self, input_path: Path, job_id: str) -> Tuple[Path, Dict[str, Any]]:
        doc = fitz.open(str(input_path))
        rendered: List[Path] = []
        page_dir = self.outputs_dir / f"{job_id}-pages"
        page_dir.mkdir(parents=True, exist_ok=True)
        for idx, page in enumerate(doc):
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            out_path = page_dir / f"page-{idx + 1}.png"
            pix.save(str(out_path))
            rendered.append(out_path)
        if len(rendered) == 1:
            preview = {"kind": "image", "url": f"/api/converter/assets/{rendered[0].name}"}
            return rendered[0], preview
        zip_path = self.outputs_dir / f"{job_id}.zip"
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
            for image_path in rendered:
                archive.write(image_path, arcname=image_path.name)
        preview = {"kind": "image", "url": f"/api/converter/assets/{rendered[0].name}", "note": "Preview shows first page. Download ZIP for all pages."}
        return zip_path, preview

    def _image_to_pdf(self, input_path: Path, output_path: Path) -> None:
        with Image.open(str(input_path)) as image:
            if image.mode in ("RGBA", "P"):
                image = image.convert("RGB")
            image.save(str(output_path), "PDF", resolution=100.0)

    def _image_to_text(self, input_path: Path, ai_mode: str, ocr_enabled: bool, structure_fix: bool, notes: List[str]) -> str:
        if not ocr_enabled:
            raise ValueError("OCR must be enabled for image to text conversion.")
        image_bytes = input_path.read_bytes()
        text, mode_note = self._ocr_image_bytes(image_bytes, ai_mode)
        if mode_note:
            notes.append(mode_note)
        return self._cleanup_text(text, ai_mode, structure_fix)

    def _ocr_image_bytes(self, image_bytes: bytes, ai_mode: str) -> Tuple[str, str]:
        tesseract_path = shutil.which("tesseract")
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
            with Image.open(io.BytesIO(image_bytes)) as image:
                return pytesseract.image_to_string(image).strip(), "OCR processed with Tesseract."

        text = self._ocr_with_ai(image_bytes, ai_mode)
        if text:
            return text, "OCR processed with AI vision fallback."
        return "", "OCR engine unavailable locally. Install Tesseract or use a vision-capable model."

    def _ocr_with_ai(self, image_bytes: bytes, ai_mode: str) -> str:
        if ai_mode == "off":
            return ""
        client, model = self._provider_client()
        if not client or not model:
            return ""
        encoded = base64.b64encode(image_bytes).decode("ascii")
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": "Extract all readable text from this document image. Preserve paragraph and heading structure. Return plain text only."},
                            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{encoded}"}},
                        ],
                    }
                ],
                temperature=0.1,
            )
            return (response.choices[0].message.content or "").strip()
        except Exception:
            return ""

    def _cleanup_text(self, text: str, ai_mode: str, structure_fix: bool) -> str:
        text = (text or "").strip()
        if not text or ai_mode == "off" or not structure_fix:
            return text
        client, model = self._provider_client()
        if not client or not model:
            return text
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {
                        "role": "system",
                        "content": "You clean OCR/document extraction text. Preserve original meaning. Do not hallucinate missing content. Return plain text only.",
                    },
                    {
                        "role": "user",
                        "content": (
                            "Normalize this extracted document text. Fix broken line wraps, restore readable paragraph flow, "
                            "and keep headings/lists where obvious.\n\n"
                            f"{text[:15000]}"
                        ),
                    },
                ],
                temperature=0.15,
            )
            cleaned = (response.choices[0].message.content or "").strip()
            return cleaned or text
        except Exception:
            return text

    def _build_source_preview(self, path: Path, ext: str) -> Dict[str, Any]:
        if ext == "txt":
            return {"kind": "text", "text": path.read_text(encoding="utf-8", errors="ignore")[:6000]}
        if ext in {"png", "jpg", "jpeg"}:
            return {"kind": "image", "url": f"/api/converter/assets/{path.name}"}
        if ext == "docx":
            try:
                return {"kind": "text", "text": self._docx_to_text(path)[:6000]}
            except Exception:
                return {"kind": "text", "text": "DOCX preview could not be extracted."}
        if ext == "pdf":
            try:
                doc = fitz.open(str(path))
                text = []
                for page in doc[:3]:
                    text.append(page.get_text("text").strip())
                rendered = self.previews_dir / f"{path.stem}-page-1.png"
                first_page = doc[0]
                first_page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False).save(str(rendered))
                return {
                    "kind": "pdf",
                    "text": "\n\n".join(t for t in text if t)[:4000],
                    "imageUrl": f"/api/converter/assets/{rendered.name}",
                }
            except Exception:
                return {"kind": "text", "text": "PDF preview is unavailable."}
        return {"kind": "text", "text": "Preview unavailable."}
